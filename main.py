"""
Hardware Pipeline — FastAPI backend.

Design principles applied here:
- Thin route handlers: parse request → call service → return response.
- No business logic in this file (lives in services/).
- CORS restricted to known origins only.
- Secrets validated at startup — missing keys cause an early, clear error.
- Pipeline runs as BackgroundTask (non-blocking, UI polls for status).
- Structured logging via Python logging throughout.
"""

import functools
import logging
import os
import pathlib
import threading
from contextlib import asynccontextmanager
from datetime import datetime
from typing import Optional

from fastapi import BackgroundTasks, FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse, RedirectResponse
from fastapi.staticfiles import StaticFiles
from starlette.middleware.base import BaseHTTPMiddleware
from pydantic import BaseModel
from typing import List

from config import settings
from logging_config import configure_logging

configure_logging()
log = logging.getLogger("hardware_pipeline.api")


# ── Startup / shutdown ────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    log.info("startup.begin", extra={"env": settings.app_env})

    # 1. Validate secrets — fail fast with a clear message
    _validate_secrets()

    # 2. Initialise DB (creates tables if they don't exist)
    from database.models import get_engine
    _db_engine = get_engine()

    # 2b. Wire up OpenTelemetry — no-op when OTEL_EXPORTER_OTLP_ENDPOINT
    # is unset, so local dev keeps working unchanged. When set, every
    # FastAPI route, SQL query, and LLM call becomes a span.
    try:
        from observability import configure_otel
        configure_otel(app=app, engine=_db_engine)
    except Exception as exc:
        log.warning("startup.otel_skipped: %s", exc)

    # 2a. Apply idempotent SQLite migrations (requirements lock columns,
    #     pipeline_runs / llm_calls tables). Safe to call on every start.
    try:
        from migrations import apply_all as _apply_migrations
        _db_url = settings.database_url
        if _db_url.startswith("sqlite:///"):
            _db_path = _db_url[len("sqlite:///"):]
            if _db_path.startswith("./"):
                _db_path = os.path.join(os.getcwd(), _db_path[2:])
            _results = _apply_migrations(_db_path)
            log.info("startup.migrations_applied", extra={"results": _results})
    except Exception as exc:
        log.warning("startup.migrations_skipped: %s", exc)
    log.info("startup.db_ready")

    # 3. Ensure output directory exists
    settings.output_dir.mkdir(parents=True, exist_ok=True)

    # 4. Seed ChromaDB component index in background — non-blocking
    def _seed_chroma():
        try:
            from tools.seed_components import seed_if_empty
            seed_if_empty()
            log.info("startup.chroma_ready")
        except Exception as exc:
            log.debug("startup.chroma_seed_skipped: %s (optional)", exc)

    threading.Thread(target=_seed_chroma, daemon=True, name="chroma-seed").start()

    log.info("startup.complete", extra={"air_gapped": settings.is_air_gapped})
    yield
    log.info("shutdown.complete")


def _validate_secrets() -> None:
    """
    Fail fast if required secrets are missing.
    Logs a clear warning for optional keys so operators know what's degraded.
    """
    if not settings.has_any_llm_key:
        if settings.app_env == "production":
            raise RuntimeError(
                "No LLM API key configured (ANTHROPIC_API_KEY or GLM_API_KEY). "
                "Set at least one before starting in production."
            )
        log.warning("startup.no_llm_key — running in air-gap/Ollama mode")

    optional_keys = {
        "DIGIKEY_CLIENT_ID": settings.digikey_client_id,
        "MOUSER_API_KEY": settings.mouser_api_key,
        "OPENAI_API_KEY": settings.openai_api_key,
    }
    for name, val in optional_keys.items():
        if not val:
            log.info("startup.optional_key_missing: %s (degraded mode)", name)


# ── App factory ────────────────────────────────────────────────────────────────

app = FastAPI(
    title=settings.app_name,
    description="AI-powered hardware design automation pipeline",
    version="2.0.0",
    lifespan=lifespan,
    docs_url="/docs" if settings.debug else None,   # hide Swagger in prod
    redoc_url=None,
)

# Static assets — mermaid.min.js and other bundled libs served from localhost
# so the frontend never needs an internet connection during the demo.
_STATIC_DIR = pathlib.Path(__file__).parent / "static"
_STATIC_DIR.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=str(_STATIC_DIR)), name="static")

# CORS: restrict to known origins only (never wildcard in any real deploy)
_ALLOWED_ORIGINS = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Content-Type", "Authorization"],
)

# ── Password Gate (optional) ───────────────────────────────────────────────────
# Set APP_PASSWORD env var to enable. Leave empty to disable (open access).
# Uses a signed cookie — no database, no sessions library needed.
# Protects all routes except /health and /login.

_APP_PASSWORD = settings.app_password or os.environ.get("APP_PASSWORD", "")
_COOKIE_NAME = "hp_auth"
_COOKIE_MAX_AGE = 60 * 60 * 24 * 7  # 7 days

def _make_token(password: str) -> str:
    """Simple HMAC token so the cookie can't be forged without knowing the password."""
    import hmac
    import hashlib
    return hmac.new(password.encode(), b"hardware-pipeline-auth", hashlib.sha256).hexdigest()

_LOGIN_PAGE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Hardware Pipeline — Login</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@700;800&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
  *, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }
  body {
    background: #070b14;
    color: #e2e8f0;
    font-family: 'DM Mono', monospace;
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    background-image: radial-gradient(circle at 50% 50%, rgba(0,198,167,0.04) 0%, transparent 70%);
  }
  .card {
    background: #1a2235;
    border: 1px solid rgba(0,198,167,0.25);
    border-radius: 12px;
    padding: 48px 44px;
    width: 100%;
    max-width: 400px;
    box-shadow: 0 0 40px rgba(0,198,167,0.08);
    text-align: center;
  }
  .logo { font-family: 'Syne', sans-serif; font-size: 26px; font-weight: 800; margin-bottom: 4px; }
  .logo span { color: #00c6a7; }
  .sub { font-size: 10px; color: #00c6a7; letter-spacing: 0.15em; margin-bottom: 32px; }
  label { display: block; font-size: 11px; color: #64748b; letter-spacing: 0.08em; margin-bottom: 8px; text-align: left; }
  input[type=password] {
    width: 100%; padding: 12px 14px;
    background: #0d1220; border: 1px solid rgba(42,58,80,0.8);
    border-radius: 6px; color: #e2e8f0; font-family: 'DM Mono', monospace;
    font-size: 14px; outline: none; margin-bottom: 18px;
    transition: border-color 0.2s;
  }
  input[type=password]:focus { border-color: #00c6a7; }
  button {
    width: 100%; padding: 12px;
    background: #00c6a7; border: none; border-radius: 6px;
    color: #070b14; font-family: 'Syne', sans-serif;
    font-size: 14px; font-weight: 700; cursor: pointer;
    letter-spacing: 0.05em; transition: opacity 0.2s;
  }
  button:hover { opacity: 0.88; }
  .err { color: #ef4444; font-size: 12px; margin-bottom: 14px; }
</style>
</head>
<body>
<div class="card">
  <div class="logo">Hardware <span>Pipeline</span></div>
  <div class="sub">DATA PATTERNS · CODE KNIGHTS</div>
  <form method="POST" action="/login">
    <label>ACCESS PASSWORD</label>
    <input type="password" name="password" placeholder="Enter password" autofocus>
    {error}
    <button type="submit">ENTER →</button>
  </form>
</div>
</body>
</html>"""

class PasswordGateMiddleware(BaseHTTPMiddleware):
    """Block all routes behind a password if APP_PASSWORD is set."""

    # Routes that bypass the gate entirely
    _OPEN = {"/health", "/login"}

    async def dispatch(self, request: Request, call_next):
        if not _APP_PASSWORD:
            return await call_next(request)  # gate disabled

        path = request.url.path
        if path in self._OPEN or path.startswith("/login"):
            return await call_next(request)

        # Check cookie
        token = request.cookies.get(_COOKIE_NAME, "")
        if token == _make_token(_APP_PASSWORD):
            return await call_next(request)

        # Not authenticated — redirect to login
        return RedirectResponse(url=f"/login?next={path}", status_code=302)

if _APP_PASSWORD:
    app.add_middleware(PasswordGateMiddleware)
    log.info("password_gate.enabled")


@app.get("/login", response_class=HTMLResponse, tags=["ops"])
async def login_page(next: str = "/app"):
    return HTMLResponse(_LOGIN_PAGE.replace("{error}", ""))


@app.post("/login", tags=["ops"])
async def login_submit(request: Request, next: str = "/app"):
    form = await request.form()
    password = form.get("password", "")
    if password == _APP_PASSWORD:
        response = RedirectResponse(url=next, status_code=302)
        response.set_cookie(
            key=_COOKIE_NAME,
            value=_make_token(_APP_PASSWORD),
            max_age=_COOKIE_MAX_AGE,
            httponly=True,
            samesite="lax",
        )
        return response
    error_html = '<div class="err">Incorrect password. Try again.</div>'
    return HTMLResponse(_LOGIN_PAGE.replace("{error}", error_html), status_code=401)


# ── Service singletons (created once per process) ─────────────────────────────
# lru_cache ensures a single instance is reused for the lifetime of the process.

@functools.lru_cache(maxsize=1)
def _project_svc():
    from services.project_service import ProjectService
    return ProjectService()

@functools.lru_cache(maxsize=1)
def _chat_svc():
    from services.chat_service import ChatService
    return ChatService()

@functools.lru_cache(maxsize=1)
def _pipeline_svc():
    from services.pipeline_service import PipelineService
    return PipelineService()


# ── Health ─────────────────────────────────────────────────────────────────────

@app.get("/health", tags=["ops"])
async def health_check():
    return {
        "status": "healthy",
        "app": settings.app_name,
        "environment": settings.app_env,
        "air_gapped": settings.is_air_gapped,
        "timestamp": datetime.utcnow().isoformat() + "Z",
    }


# ── Projects ───────────────────────────────────────────────────────────────────

@app.post("/api/v1/projects", status_code=201, tags=["projects"])
async def create_project(body: dict):
    name = (body.get("name") or "").strip()
    if not name:
        raise HTTPException(400, "name is required")
    design_scope = (body.get("design_scope") or "full").strip().lower()
    project_type = (body.get("project_type") or "receiver").strip().lower()
    if project_type not in ("receiver", "transmitter"):
        raise HTTPException(400, f"project_type must be 'receiver' or 'transmitter' (got '{project_type}')")
    try:
        return _project_svc().create(
            name=name,
            description=body.get("description", ""),
            design_type=body.get("design_type", "rf"),
            design_scope=design_scope,
            project_type=project_type,
        )
    except ValueError as exc:
        raise HTTPException(400, str(exc))
    except Exception as exc:
        log.exception("api.create_project_failed")
        raise HTTPException(500, str(exc))


class DesignScopeBody(BaseModel):
    design_scope: str

@app.patch("/api/v1/projects/{project_id}/design-scope", tags=["projects"])
async def update_design_scope(project_id: int, body: DesignScopeBody):
    """Update the wizard-selected design scope after creation.

    The wizard may narrow or widen the scope mid-session; the frontend
    persists that choice here so the backend can enforce `applicableScopes`
    at phase-execute time without trusting localStorage."""
    try:
        return _project_svc().set_design_scope(project_id, body.design_scope)
    except ValueError as exc:
        msg = str(exc)
        if "not found" in msg.lower():
            raise HTTPException(404, msg)
        raise HTTPException(400, msg)


@app.get("/api/v1/projects", tags=["projects"])
async def list_projects():
    return _project_svc().list_all()


@app.get("/api/v1/projects/{project_id}", tags=["projects"])
async def get_project(project_id: int):
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    return proj


def _resolve_output_dir(proj: dict) -> Optional[str]:
    """
    Return the output directory for a project, with fallback derivation.

    Priority:
    1. DB-stored output_dir (absolute or relative path that exists on disk)
    2. Derived from project name using the same StorageAdapter logic
       (handles projects created before output_dir was reliably written, or
        where the DB column was left empty due to a failed project creation)
    """
    stored = (proj.get("output_dir") or "").strip()
    if stored and os.path.isdir(stored):
        return stored

    # Fallback: derive from project name using StorageAdapter.project_dir logic
    name = (proj.get("name") or "").strip()
    if name:
        safe = name.replace(" ", "_").lower()
        # Try relative (server started from project root) and absolute via settings
        candidates = [
            os.path.join("output", safe),
            str(settings.output_dir / safe),
        ]
        for candidate in candidates:
            if os.path.isdir(candidate):
                log.info(
                    "documents.output_dir_derived",
                    extra={"project_id": proj.get("id"), "derived": candidate, "stored": stored or "(empty)"},
                )
                return candidate

    log.warning(
        "documents.output_dir_missing",
        extra={"project_id": proj.get("id"), "stored": stored or "(empty)", "name": proj.get("name")},
    )
    return None


@app.get("/api/v1/projects/{project_id}/documents/{filename:path}", tags=["projects"])
async def get_document(project_id: int, filename: str):
    # :path type captures slashes, so qt_gui/ControlPanel.cpp works as-is
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "Project output directory not found — run Phase 1 first")

    # Guard against path traversal
    base = os.path.realpath(output_dir)
    file_path = os.path.realpath(os.path.join(output_dir, filename))
    if not file_path.startswith(base):
        raise HTTPException(400, "Invalid filename")
    if not os.path.exists(file_path):
        raise HTTPException(404, f"Document {filename} not found")

    return FileResponse(file_path)


@app.get("/api/v1/projects/{project_id}/documents", tags=["projects"])
async def list_documents(project_id: int):
    """List all available output files for a project (flat + one level deep for qt_gui etc.)."""
    proj = _project_svc().get(project_id)
    if not proj:
        return []

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        return []

    try:
        _HIDDEN_DIRS = {".docx_cache"}
        files = []
        for entry in os.scandir(output_dir):
            if entry.is_file():
                files.append({"name": entry.name, "size": entry.stat().st_size})
            elif entry.is_dir() and entry.name not in _HIDDEN_DIRS:
                # Include one level of subdirectory files (e.g. qt_gui/, .github/workflows/)
                try:
                    for sub in os.scandir(entry.path):
                        if sub.is_file():
                            rel = f"{entry.name}/{sub.name}"
                            files.append({"name": rel, "size": sub.stat().st_size})
                except OSError:
                    pass
        return sorted(files, key=lambda x: x["name"])
    except OSError as exc:
        log.warning("documents.list_failed", extra={"project_id": project_id, "error": str(exc)})
        return []


# ── Clarify (Phase 1 — pre-chat card generation) ──────────────────────────────

class ClarifyRequest(BaseModel):
    requirement: str
    design_type: str = "RF"
    # Optional prior turns so later elicitation rounds can reuse /clarify for
    # structured follow-up cards. Each entry: { role: "user"|"assistant", content: str }.
    conversation_history: list[dict] | None = None
    # Round label for logging / prompt hinting (e.g. "round-2-architecture").
    round_label: str | None = None


@app.post("/api/v1/projects/{project_id}/clarify", tags=["chat"])
async def get_clarification_questions(project_id: int, body: ClarifyRequest):
    """
    Return structured clarification questions as interactive card data —
    powered by tool_use (zero parse failures).

    * First call (round-1): no conversation_history, just a raw requirement.
    * Follow-up calls (round-2+): include conversation_history so the agent can
      see the prior turns and produce the NEXT round of cards without repeating
      earlier questions. Frontend renders these as clickable option cards, then
      bundles the answers into one message and forwards it to POST /chat.
    """
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    try:
        from agents.requirements_agent import RequirementsAgent
        agent = RequirementsAgent()
        result = agent.get_clarification_questions(
            user_requirement=body.requirement,
            design_type=body.design_type,
            conversation_history=body.conversation_history,
            round_label=body.round_label,
        )
        return result
    except ValueError as exc:
        raise HTTPException(502, str(exc))
    except Exception as exc:
        log.exception("api.clarify_failed", extra={"project_id": project_id})
        raise HTTPException(500, f"Clarification failed: {str(exc)}")


# ── Chat (Phase 1) ─────────────────────────────────────────────────────────────

# Wall-clock cap on one chat turn. The inner LLM call already has its
# own per-request timeout, but without an outer ceiling a slow/hung LLM
# or stalled distributor lookup would keep the HTTP connection open
# indefinitely. Default 600 s covers a heavy P1 finalize (LLM + cascade +
# audit + distributor validation) traversing the full fallback chain
# (glm-4.7 → glm-4.5-air → deepseek-chat), where each hop can burn
# ~120 s on its own httpx timeout before handing off. Override with
# CHAT_DEADLINE_S in .env if the upstream LLM is unusually slow.
_CHAT_DEADLINE_S = float(os.environ.get("CHAT_DEADLINE_S", "600"))


@app.post("/api/v1/projects/{project_id}/chat", tags=["chat"])
async def chat(project_id: int, body: dict):
    """Send a message to the Phase 1 requirements agent."""
    import asyncio as _asyncio
    message = (body.get("message") or "").strip()
    if not message:
        raise HTTPException(400, "message is required")

    try:
        result = await _asyncio.wait_for(
            _chat_svc().send_message(project_id, message),
            timeout=_CHAT_DEADLINE_S,
        )
        log.info("api.chat_ok",
                 extra={"project_id": project_id, "phase_complete": result.get("phase_complete")})
        return result
    except _asyncio.TimeoutError:
        log.warning(
            "api.chat_timeout project_id=%s deadline=%.0fs",
            project_id, _CHAT_DEADLINE_S,
        )
        raise HTTPException(
            504,
            f"Chat request exceeded the {_CHAT_DEADLINE_S:.0f}s deadline. "
            "Try again, or shorten the request.",
        )
    except ValueError as exc:
        raise HTTPException(404, str(exc))
    except Exception as exc:
        log.exception("api.chat_failed", extra={"project_id": project_id})
        raise HTTPException(500, str(exc))


# ── Configuration Settings ────────────────────────────────────────────────────────

_ENV_FILE = pathlib.Path(__file__).parent / ".env"

def _persist_env(updates: dict) -> None:
    """Write key=value pairs to .env file, creating or updating entries."""
    try:
        lines = _ENV_FILE.read_text(encoding="utf-8").splitlines() if _ENV_FILE.exists() else []
        updated_keys = set()
        new_lines = []
        for line in lines:
            stripped = line.strip()
            if stripped and not stripped.startswith("#"):
                key = stripped.split("=", 1)[0].strip()
                if key in updates:
                    new_lines.append(f"{key}={updates[key]}")
                    updated_keys.add(key)
                    continue
            new_lines.append(line)
        # Append any keys not yet in the file
        for key, val in updates.items():
            if key not in updated_keys:
                new_lines.append(f"{key}={val}")
        _ENV_FILE.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
    except Exception as exc:
        log.warning("settings.persist_env_failed: %s", exc)


class ConfigSettingsRequest(BaseModel):
    # LLM keys
    glm_api_key: Optional[str] = None
    deepseek_api_key: Optional[str] = None
    anthropic_api_key: Optional[str] = None
    glm_base_url: Optional[str] = None
    deepseek_base_url: Optional[str] = None
    primary_model: Optional[str] = None
    fast_model: Optional[str] = None
    # GitHub / Git
    github_token: Optional[str] = None
    github_repo: Optional[str] = None
    # Component search
    digikey_client_id: Optional[str] = None
    digikey_client_secret: Optional[str] = None
    mouser_api_key: Optional[str] = None
    chroma_persist_dir: Optional[str] = None


# Keep old name as alias for backward compat
LLMSettingsRequest = ConfigSettingsRequest


def _mask_key(key: Optional[str]) -> Optional[str]:
    if not key:
        return None
    if len(key) <= 8:
        return "•" * len(key)
    return key[:6] + "•" * min(len(key) - 10, 12) + key[-4:]


@app.get("/api/v1/settings/llm", tags=["settings"])
async def get_llm_settings():
    """Get current configuration. Returns masked API keys for security."""
    return {
        "glm_api_key": _mask_key(settings.glm_api_key),
        "deepseek_api_key": _mask_key(settings.deepseek_api_key),
        "anthropic_api_key": _mask_key(settings.anthropic_api_key),
        "glm_base_url": settings.glm_base_url,
        "deepseek_base_url": settings.deepseek_base_url,
        "primary_model": settings.primary_model,
        "fast_model": settings.fast_model,
        "glm_model": settings.glm_model,
        "glm_fast_model": settings.glm_fast_model,
        "github_token": _mask_key(settings.github_token),
        "github_repo": settings.github_repo,
        "git_enabled": settings.git_enabled,
        "digikey_client_id": _mask_key(settings.digikey_client_id),
        "digikey_client_secret": _mask_key(settings.digikey_client_secret),
        "mouser_api_key": _mask_key(settings.mouser_api_key),
        "chroma_persist_dir": settings.chroma_persist_dir,
    }


@app.post("/api/v1/settings/llm", tags=["settings"])
async def update_llm_settings(body: ConfigSettingsRequest):
    """
    Update LLM + GitHub configuration. Persists changes to .env file.
    """
    env_updates: dict = {}

    def _apply(attr: str, env_key: str, value: Optional[str]) -> None:
        if value is not None and value.strip():
            v = value.strip()
            setattr(settings, attr, v)
            os.environ[env_key] = v
            env_updates[env_key] = v

    _apply("glm_api_key",      "GLM_API_KEY",       body.glm_api_key)
    _apply("deepseek_api_key", "DEEPSEEK_API_KEY",   body.deepseek_api_key)
    _apply("anthropic_api_key","ANTHROPIC_API_KEY",  body.anthropic_api_key)
    _apply("glm_base_url",     "GLM_BASE_URL",       body.glm_base_url)
    _apply("deepseek_base_url","DEEPSEEK_BASE_URL",  body.deepseek_base_url)
    _apply("primary_model",    "PRIMARY_MODEL",      body.primary_model)
    _apply("fast_model",       "FAST_MODEL",         body.fast_model)
    _apply("github_token",          "GITHUB_TOKEN",          body.github_token)
    _apply("github_repo",           "GITHUB_REPO",           body.github_repo)
    _apply("digikey_client_id",     "DIGIKEY_CLIENT_ID",     body.digikey_client_id)
    _apply("digikey_client_secret", "DIGIKEY_CLIENT_SECRET", body.digikey_client_secret)
    _apply("mouser_api_key",        "MOUSER_API_KEY",        body.mouser_api_key)
    _apply("chroma_persist_dir",    "CHROMA_PERSIST_DIR",    body.chroma_persist_dir)

    # Recompute git_enabled after token change
    if body.github_token is not None:
        settings.git_enabled = bool(settings.github_token)
        os.environ["GIT_ENABLED"] = "true" if settings.git_enabled else "false"

    # Persist to .env so settings survive restarts
    if env_updates:
        _persist_env(env_updates)

    log.info("api.settings_updated", extra={
        "keys_updated": list(env_updates.keys()),
        "git_enabled": settings.git_enabled,
    })

    return {
        "glm_api_key": _mask_key(settings.glm_api_key),
        "deepseek_api_key": _mask_key(settings.deepseek_api_key),
        "anthropic_api_key": _mask_key(settings.anthropic_api_key),
        "glm_base_url": settings.glm_base_url,
        "deepseek_base_url": settings.deepseek_base_url,
        "primary_model": settings.primary_model,
        "fast_model": settings.fast_model,
        "glm_model": settings.glm_model,
        "glm_fast_model": settings.glm_fast_model,
        "github_token": _mask_key(settings.github_token),
        "github_repo": settings.github_repo,
        "git_enabled": settings.git_enabled,
        "digikey_client_id": _mask_key(settings.digikey_client_id),
        "digikey_client_secret": _mask_key(settings.digikey_client_secret),
        "mouser_api_key": _mask_key(settings.mouser_api_key),
        "chroma_persist_dir": settings.chroma_persist_dir,
    }


# ── Pipeline (P2→P8c background execution) ────────────────────────────────────

@app.post("/api/v1/projects/{project_id}/pipeline/run", tags=["pipeline"])
async def run_pipeline(project_id: int, background_tasks: BackgroundTasks):
    """
    Start the full P2→P8c pipeline as a background task.
    Returns immediately; UI should poll GET /projects/{id} for status.
    """
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    p1_status = _project_svc().get_phase_status(project_id, "P1")
    if p1_status not in ("completed", "draft_pending"):
        raise HTTPException(400, "Phase 1 must be completed before running the pipeline")

    # User clicked "Approve & Start Pipeline" — if P1 is in draft_pending, promote to completed
    if p1_status == "draft_pending":
        _project_svc().set_phase_status(project_id, "P1", "completed")
        log.info("api.p1_approved_and_completed", extra={"project_id": project_id})

    svc = _pipeline_svc()
    background_tasks.add_task(svc.run_pipeline, project_id)
    log.info("api.pipeline_started", extra={"project_id": project_id})
    return {"status": "pipeline_started", "project_id": project_id}


VALID_PHASES = {"P1", "P2", "P3", "P4", "P5", "P6", "P7", "P7a", "P8a", "P8b", "P8c"}

@app.post("/api/v1/projects/{project_id}/phases/{phase_id}/execute", tags=["pipeline"])
async def execute_single_phase(project_id: int, phase_id: str, background_tasks: BackgroundTasks):
    """Execute one specific phase as a background task.

    Enforces the project's `design_scope` — returns 409 Conflict if the
    requested phase is not applicable for the project's wizard-selected scope.
    This is what stops an out-of-scope phase from silently running despite
    the UI labelling it NOT APPLICABLE.
    """
    if phase_id not in VALID_PHASES:
        raise HTTPException(400, f"Invalid phase '{phase_id}'. Must be one of: {sorted(VALID_PHASES)}")
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    # Scope gate — backend is now the source of truth (not localStorage).
    from services.phase_scopes import is_phase_applicable
    scope = (proj.get("design_scope") or "full").lower()
    if not is_phase_applicable(phase_id, scope):
        log.warning(
            "api.phase_not_applicable",
            extra={"project_id": project_id, "phase_id": phase_id, "design_scope": scope},
        )
        raise HTTPException(
            409,
            f"Phase {phase_id} is not applicable for design_scope '{scope}'. "
            "Change the project's scope or pick a different phase.",
        )

    try:
        svc = _pipeline_svc()
        background_tasks.add_task(svc.run_single_phase, project_id, phase_id)
        return {"status": "phase_started", "phase_id": phase_id, "project_id": project_id}
    except ValueError as exc:
        raise HTTPException(400, str(exc))


@app.post("/api/v1/projects/{project_id}/phases/{phase_id}/cancel", tags=["pipeline"])
async def cancel_phase(project_id: int, phase_id: str):
    """Cancel a running phase by setting its status to 'pending'.
    Note: the background LLM task may still complete, but the frontend
    will stop polling and the phase can be re-run cleanly."""
    if phase_id not in VALID_PHASES:
        raise HTTPException(400, f"Invalid phase '{phase_id}'")
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    _project_svc().set_phase_status(project_id, phase_id, "pending")
    log.info("api.phase_cancelled", extra={"project_id": project_id, "phase_id": phase_id})
    return {"status": "cancelled", "phase_id": phase_id, "project_id": project_id}


class ResetPhasesRequest(BaseModel):
    phase_ids: List[str]

@app.post("/api/v1/projects/{project_id}/phases/reset", tags=["pipeline"])
async def reset_phases(project_id: int, body: ResetPhasesRequest, background_tasks: BackgroundTasks):
    """
    Reset given phases to 'pending' then immediately re-run the pipeline.
    Used by the frontend 'Re-run all stale' button after requirements change.
    """
    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    if not body.phase_ids:
        raise HTTPException(400, "phase_ids must not be empty")

    # Validate phase IDs
    invalid = [p for p in body.phase_ids if p not in VALID_PHASES]
    if invalid:
        raise HTTPException(400, f"Invalid phase IDs: {invalid}")

    # Reset each phase status to pending
    svc = _project_svc()
    for phase_id in body.phase_ids:
        svc.set_phase_status(project_id, phase_id, "pending")

    log.info("api.phases_reset", extra={"project_id": project_id, "phase_ids": body.phase_ids})

    # Kick off the pipeline — it will now run all the reset phases
    pipeline = _pipeline_svc()
    background_tasks.add_task(pipeline.run_pipeline, project_id)
    return {"status": "pipeline_started", "reset_phases": body.phase_ids, "project_id": project_id}


@app.get("/api/v1/projects/{project_id}/export", tags=["projects"])
async def export_project_zip(project_id: int):
    """
    Stream all project output documents as a ZIP archive.
    Frontend uses this for the 'Download All Documents' button.
    """
    import io
    import zipfile
    import pathlib

    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "No output directory for this project")

    out_path = pathlib.Path(output_dir)
    if not out_path.exists():
        raise HTTPException(404, "Output directory does not exist")

    doc_files = [f for f in out_path.rglob("*") if f.is_file() and ".docx_cache" not in f.parts]
    if not doc_files:
        raise HTTPException(404, "No documents found to export")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in doc_files:
            zf.write(f, f.relative_to(out_path))
    buf.seek(0)

    safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in (proj.get("name") or "project"))
    filename = f"{safe_name}_documents.zip"

    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _render_mermaid_local(code: str, out_path: str) -> bool:
    """
    Render a Mermaid diagram to PNG.

    1. mermaid.ink — public REST API, fast, no local deps (requires internet).
    2. mmdc       — @mermaid-js/mermaid-cli (local fallback).
    3. Node.js    — bundled renderer + cairosvg (last resort).

    Returns True on success, False when all methods fail.
    """
    import subprocess as _sp
    import pathlib as _pl
    import tempfile as _tmp
    import os as _os

    # ── 1. mermaid.ink public API — fastest, no local tools ──────────────────
    try:
        import base64 as _b64, urllib.request as _req
        encoded = _b64.urlsafe_b64encode(code.encode('utf-8')).decode()
        url = f"https://mermaid.ink/img/{encoded}?type=png&bgColor=white&width=1400"
        req = _req.Request(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'image/png'
        })
        with _req.urlopen(req, timeout=15) as resp:
            data = resp.read()
        with open(out_path, 'wb') as f:
            f.write(data)
        if _pl.Path(out_path).exists() and _pl.Path(out_path).stat().st_size > 200:
            log.debug("mermaid.ink.ok")
            return True
    except Exception as _ink_e:
        log.debug("mermaid.ink.skip: %s", _ink_e)

    # ── 2. mmdc (mermaid-cli) — local fallback ───────────────────────────────
    try:
        tmp_dir_obj = _tmp.TemporaryDirectory()
        tmp_dir = tmp_dir_obj.name
        mmd_file = _pl.Path(tmp_dir) / 'diagram.mmd'
        mmd_file.write_text(code, encoding='utf-8')

        mmdc_cmd = 'mmdc.cmd' if _os.name == 'nt' else 'mmdc'
        result = _sp.run(
            [mmdc_cmd, '-i', str(mmd_file), '-o', out_path,
             '-b', 'white', '-w', '1400', '--quiet'],
            capture_output=True, timeout=15,
        )

        try:
            tmp_dir_obj.cleanup()
        except Exception:
            pass

        if result.returncode == 0 and _pl.Path(out_path).exists():
            if _pl.Path(out_path).stat().st_size > 200:
                log.debug("mermaid.mmdc.ok")
                return True
    except (FileNotFoundError, _sp.TimeoutExpired, Exception) as _e:
        log.debug("mermaid.mmdc.skip: %s", _e)

    # ── 3. Bundled Node.js renderer (last resort) ────────────────────────────
    _renderer_candidates = [
        _pl.Path(__file__).parent / "mermaid_renderer.js",
        _pl.Path(__file__).parent.parent / "mermaid-renderer" / "render.js",
    ]
    _renderer_js = next((str(p) for p in _renderer_candidates if p.exists()), None)
    if _renderer_js:
        try:
            result = _sp.run(
                ["node", _renderer_js, code],
                capture_output=True, text=True, timeout=15,
            )
            svg_str = result.stdout
            if result.returncode == 0 and svg_str and "<svg" in svg_str:
                try:
                    from cairosvg import svg2png  # type: ignore
                    png_data = svg2png(
                        bytestring=svg_str.encode("utf-8"),
                        scale=2.0, background_color="white",
                    )
                    if png_data and len(png_data) > 200:
                        _pl.Path(out_path).write_bytes(png_data)
                        log.debug("mermaid.node.ok: %s", _renderer_js)
                        return True
                except Exception as _cairo_e:
                    log.debug("mermaid.cairosvg.skip: %s", _cairo_e)
            else:
                log.debug("mermaid.node.skip: %s", result.stderr[:200])
        except Exception as _node_e:
            log.debug("mermaid.node.error: %s", _node_e)

    log.warning("mermaid.render.all_failed")
    return False


def _sanitize_mermaid_code(code: str) -> str:
    """
    Sanitize mermaid code to fix common AI-generation errors.
    This must match the frontend sanitization in DocumentsView.tsx and ChatView.tsx.
    """
    import re as _re

    # Strip %%{ init }%% frontmatter and %% comments
    code = _re.sub(r'^%%\{[\s\S]*?\}%%\s*', '', code, flags=_re.MULTILINE)
    code = _re.sub(r'%%[^\n]*', '', code)

    # Arrow fixes: em-dash → --, ——> → -->
    code = code.replace('\u2014\u2014>', '-->').replace('\u2014>', '-->')
    code = code.replace('——>', '-->').replace('—>', '-->')
    code = code.replace('==>', '-->')
    code = _re.sub(r'(\w)\s*->\s*(\w)', r'\1 --> \2', code)

    # Single dash between nodes → -->
    code = _re.sub(r'(\b[\w\-]+\b)\s+-\s+(\b[\w\-]+[\s\[\(])', r'\1 --> \2', code)

    # Normalise graph → flowchart
    code = _re.sub(r'^graph\s+(TD|LR|TB|RL|BT)', r'flowchart \1', code, flags=_re.IGNORECASE)
    code = _re.sub(r'^(flowchart)\n(TD|LR|TB|RL|BT)\b', r'\1 \2', code, flags=_re.MULTILINE)

    # Join multi-line labels
    lines = code.split('\n')
    joined = []
    for line in lines:
        if joined and _re.search(r'\[', joined[-1]) and not _re.search(r'\]', joined[-1]):
            joined[-1] = joined[-1].rstrip() + ' ' + line.lstrip()
        else:
            joined.append(line)
    code = '\n'.join(joined)

    # Strip non-ASCII characters that break Mermaid (Ω, °, µ, etc.)
    code = _re.sub(r'[^\x00-\x7F]', lambda m: {
        '\u03a9': 'Ohm', '\u00b0': 'deg', '\u00b5': 'u', '\u2126': 'Ohm',
        '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2264': '<=', '\u2265': '>=',
    }.get(m.group(), ''), code)

    # Sanitize node labels
    def sanitize_label(inner: str) -> str:
        s = inner
        s = _re.sub(r'-->', ' ', s)
        s = _re.sub(r'->', ' ', s)
        s = s.replace('<', ' ').replace('>', ' ')
        s = s.replace('(', ' ').replace(')', ' ')
        s = s.replace('_', '-')
        s = _re.sub(r'&(?!amp;|lt;|gt;|#)', 'and', s)
        s = s.replace('"', ' ').replace("'", ' ')
        s = s.replace('#', ' ')
        s = s.replace('|', '/')
        # IMPORTANT: Remove ALL dash sequences (2 or more)
        s = _re.sub(r'-{2,}', ' ', s)
        s = _re.sub(r'^[-—=]+|[—=-]+$', ' ', s)
        s = _re.sub(r'\s{2,}', ' ', s)
        return s.strip()

    # Apply to square brackets, round parens, braces (node labels)
    code = _re.sub(r'\[([^\]]*)\]', lambda m: f'[{sanitize_label(m.group(1))}]', code)
    code = _re.sub(r'\(([^)]*)\)', lambda m: f'({sanitize_label(m.group(1))})', code)
    code = _re.sub(r'\{([^}]*)\}', lambda m: f'{{{sanitize_label(m.group(1))}}}', code)

    # Sanitize edge labels: -->|label text| — strip problematic chars inside pipes
    def sanitize_edge_label(m):
        arrow = m.group(1)
        label = m.group(2)
        label = label.replace('"', ' ').replace("'", ' ')
        label = label.replace('(', ' ').replace(')', ' ')
        label = label.replace('<', ' ').replace('>', ' ')
        label = label.replace('#', ' ')
        label = _re.sub(r'[^\x20-\x7E]', '', label)  # ASCII printable only
        label = _re.sub(r'\s{2,}', ' ', label).strip()
        return f'{arrow}|{label}|'
    code = _re.sub(r'(-->|---|-\.-|==>)\|([^|]*)\|', sanitize_edge_label, code)

    # Fix double-paren nodes like ((label)) — replace with single-paren (label)
    # Mermaid ((...)) is stadium shape but AI often generates broken nested parens
    code = _re.sub(r'\(\(([^)]*)\)\)', r'(\1)', code)

    # Fix stray unmatched parens/brackets in node IDs (not inside labels)
    # e.g. "A))" at end of line → "A"
    lines = code.split('\n')
    cleaned = []
    for line in lines:
        # Remove trailing )) or ]] not part of a label
        if not _re.search(r'[\[\(\{]', line):
            line = _re.sub(r'[)\]]+\s*$', '', line)
        cleaned.append(line)
    code = '\n'.join(cleaned)

    return code


def _render_mermaid_diagrams_sync(md_text: str, tmp_dir: str) -> str:
    """
    Pre-render ```mermaid``` blocks to PNG images using local Node.js + cairosvg.
    All diagrams are rendered IN PARALLEL (ThreadPoolExecutor).
    Failures fall back gracefully to a labelled code block.
    """
    import re as _re
    import pathlib as _pl
    from concurrent.futures import ThreadPoolExecutor, as_completed

    MERMAID_RE = _re.compile(r'```mermaid\s*\n([\s\S]*?)```', _re.IGNORECASE)
    tmp = _pl.Path(tmp_dir)

    # ── 1. Collect all mermaid blocks ─────────────────────────────────────────
    blocks = []  # list of (match, code)
    for m in MERMAID_RE.finditer(md_text):
        # Sanitize the mermaid code BEFORE rendering to fix AI-generation errors
        original_code = m.group(1).strip()
        sanitized_code = _sanitize_mermaid_code(original_code)
        blocks.append((m, sanitized_code))

    if not blocks:
        return md_text

    # ── 2. Render all diagrams in parallel ────────────────────────────────────
    def render_diagram(idx_code):
        idx, code = idx_code
        img_path = str(tmp / f"diagram_{idx}.png")
        log.debug(f"mermaid.render.start idx={idx} path={img_path}")
        success = _render_mermaid_local(code, img_path)
        if success:
            log.info(f"mermaid.render.ok idx={idx} size={_pl.Path(img_path).stat().st_size if _pl.Path(img_path).exists() else 0}")
        else:
            log.warning(f"mermaid.render.failed idx={idx}")
        return idx, img_path if success else None

    results: dict[int, str | None] = {}
    with ThreadPoolExecutor(max_workers=min(len(blocks), 4)) as pool:
        futures = {pool.submit(render_diagram, (i + 1, code)): i for i, (_, code) in enumerate(blocks)}
        for fut in as_completed(futures):
            idx, path = fut.result()
            results[idx] = path

    log.info(f"mermaid.render.summary total={len(blocks)} success={sum(1 for p in results.values() if p)} failed={sum(1 for p in results.values() if not p)}")

    # ── 3. Replace blocks in reverse order (preserves string offsets) ─────────
    result_md = md_text
    for i, (m, code) in reversed(list(enumerate(blocks))):
        idx = i + 1
        img_path = results.get(idx)
        if img_path:
            # Convert Windows backslashes to forward slashes for markdown compatibility
            # Pandoc and python-docx work with forward slashes even on Windows
            img_path_md = img_path.replace('\\', '/')
            replacement = f"\n\n**System Architecture Diagram {idx}**\n\n![Diagram {idx}]({img_path_md})\n\n"
        else:
            replacement = (
                f"\n\n**System Architecture Diagram {idx}** "
                f"*(rendered in browser — source below)*\n\n"
                f"```\n{code}\n```\n\n"
            )
        result_md = result_md[:m.start()] + replacement + result_md[m.end():]

    return result_md


@app.get("/api/v1/projects/{project_id}/docx/{filename:path}", tags=["projects"])
async def convert_document_to_docx(project_id: int, filename: str):
    """
    Convert a Markdown (.md) file to .docx and stream it for download.
    Mermaid diagrams are pre-rendered to PNG via mermaid.ink before conversion.
    Uses pandoc if available, falls back to python-docx.
    Converted .docx files are cached on disk next to the source .md file.
    """
    import subprocess
    import tempfile
    import pathlib

    proj = _project_svc().get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    output_dir = _resolve_output_dir(proj)
    if not output_dir:
        raise HTTPException(404, "Project output directory not found")

    src_path = pathlib.Path(output_dir) / filename
    if not src_path.exists():
        raise HTTPException(404, f"File {filename} not found")

    if src_path.suffix.lower() not in (".md", ".txt"):
        raise HTTPException(400, "Only .md and .txt files can be converted to .docx")

    stem = src_path.stem
    out_filename = f"{stem}.docx"

    # ── Disk cache: serve cached .docx if source hasn't changed ───────────────
    # Cache lives in a hidden sub-folder so it never appears in document lists
    # and old pre-mermaid cached files are automatically bypassed.
    cache_dir = src_path.parent / ".docx_cache"
    cache_dir.mkdir(exist_ok=True)
    cache_path = cache_dir / out_filename
    if cache_path.exists() and cache_path.stat().st_mtime >= src_path.stat().st_mtime:
        log.info("docx.cache_hit", extra={"file": filename})
        cached_data = cache_path.read_bytes()
        return StreamingResponse(
            iter([cached_data]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
        )

    # ── Try pandoc first (installed in Docker image) ───────────────────────────
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Pre-render Mermaid blocks → PNG images in a thread (non-blocking)
            import asyncio
            import functools
            raw_md = src_path.read_text(encoding="utf-8")
            loop = asyncio.get_event_loop()
            processed_md = await loop.run_in_executor(
                None, functools.partial(_render_mermaid_diagrams_sync, raw_md, tmpdir)
            )

            # Write the processed markdown to a temp file in the same tmpdir
            # (so relative image paths resolve correctly for pandoc)
            tmp_md = pathlib.Path(tmpdir) / f"{stem}_processed.md"
            tmp_md.write_text(processed_md, encoding="utf-8")

            out_path = pathlib.Path(tmpdir) / out_filename
            result = subprocess.run(
                ["pandoc", str(tmp_md), "-o", str(out_path),
                 "--from=markdown", "--to=docx",
                 "-V", "geometry:margin=2.5cm",
                 "--resource-path", str(tmpdir),  # Tell pandoc where to find images
                 "--standalone"],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode == 0 and out_path.exists():
                data = out_path.read_bytes()
                # Write to disk cache for future requests
                try:
                    cache_path.write_bytes(data)
                except Exception as cache_err:
                    log.warning("docx.cache_write_failed", extra={"error": str(cache_err)})
                return StreamingResponse(
                    iter([data]),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
                )
            log.warning("pandoc.failed", extra={"stderr": result.stderr, "file": filename})
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError, Exception) as exc:
        log.warning("pandoc.unavailable", extra={"error": str(exc)})

    # ── Fallback: python-docx full markdown parser ────────────────────────────
    try:
        from docx import Document as DocxDocument  # type: ignore
        from docx.shared import Pt, RGBColor, Inches  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        import io as _io

        raw_md_text = src_path.read_text(encoding="utf-8")

        # Pre-render mermaid blocks → PNG images (same as pandoc path)
        import asyncio
        import functools
        import tempfile as _tempfile
        _fallback_tmpdir_obj = _tempfile.TemporaryDirectory()
        _fallback_tmpdir = _fallback_tmpdir_obj.name
        try:
            loop2 = asyncio.get_event_loop()
            md_text = await loop2.run_in_executor(
                None, functools.partial(_render_mermaid_diagrams_sync, raw_md_text, _fallback_tmpdir)
            )
        except Exception as _merm_exc:
            log.warning("docx.fallback.mermaid_failed", extra={"error": str(_merm_exc)})
            md_text = raw_md_text

        doc = DocxDocument()

        # Style the default Normal paragraph
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        lines = md_text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            s = line.strip()

            # ── Headings ──────────────────────────────────────────────────────
            if s.startswith("#### "):
                doc.add_heading(s[5:], level=4)
            elif s.startswith("### "):
                doc.add_heading(s[4:], level=3)
            elif s.startswith("## "):
                doc.add_heading(s[3:], level=2)
            elif s.startswith("# "):
                doc.add_heading(s[2:], level=1)

            # ── Table: collect all pipe-rows ──────────────────────────────────
            elif s.startswith("|") and s.endswith("|"):
                # Gather contiguous table rows
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                # Parse header row (first line) and separator (second line)
                def _split_row(row: str):
                    parts = [c.strip() for c in row.strip('|').split('|')]
                    return parts
                header_cells = _split_row(table_lines[0])
                # Find data rows (skip separator row which only has --- chars)
                data_rows = [_split_row(r) for r in table_lines[1:]
                             if not all(c.strip('-:') == '' for c in _split_row(r))]
                num_cols = len(header_cells)
                tbl = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                tbl.style = 'Table Grid'
                # Header row
                hdr = tbl.rows[0]
                for ci, cell_text in enumerate(header_cells[:num_cols]):
                    cell = hdr.cells[ci]
                    cell.text = cell_text
                    cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else None
                    p = cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].bold = True
                # Data rows
                for ri, row_cells in enumerate(data_rows):
                    row = tbl.rows[ri + 1]
                    for ci, cell_text in enumerate(row_cells[:num_cols]):
                        row.cells[ci].text = cell_text
                doc.add_paragraph("")  # spacing after table
                continue  # i was already advanced inside the while loop

            # ── Bullet list ───────────────────────────────────────────────────
            elif s.startswith("- ") or s.startswith("* "):
                doc.add_paragraph(s[2:], style='List Bullet')
            elif len(s) > 1 and s[0].isdigit() and s[1] in '.):':
                doc.add_paragraph(s[2:].strip(), style='List Number')

            # ── Inline image: ![alt](path) ────────────────────────────────────
            elif s.startswith("![") and "](" in s and s.endswith(")"):
                import re as _re2
                img_m = _re2.match(r'!\[([^\]]*)\]\(([^)]+)\)', s)
                if img_m:
                    img_alt, img_path = img_m.group(1), img_m.group(2).strip()
                    import pathlib as _pl2
                    _img_p = _pl2.Path(img_path)
                    log.debug(f"docx.img path={_img_p} exists={_img_p.exists()}")
                    if _img_p.exists() and _img_p.suffix.lower() in (".png", ".jpg", ".jpeg", ".gif"):
                        try:
                            # Get image size to determine width
                            from PIL import Image as _PILImage
                            with _PILImage.open(_img_p) as _img:
                                _width, _height = _img.size
                                # Scale to max 6 inches wide
                                _scaled_width = min(_width / 100.0, 6.0)
                            doc.add_picture(str(_img_p), width=Inches(_scaled_width))
                            if img_alt and img_alt != f"Diagram {img_alt}":
                                cap = doc.add_paragraph(img_alt)
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as _img_err:
                            log.warning(f"docx.img.error path={_img_p} error={_img_err}")
                            doc.add_paragraph(f"[Diagram: {img_alt}]")
                    else:
                        log.warning(f"docx.img.not_found path={_img_p}")
                        doc.add_paragraph(f"[Diagram: {img_alt}]")

            # ── Code block (skip non-mermaid fenced blocks) ───────────────────
            elif s.startswith("```"):
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    i += 1

            # ── Horizontal rule ───────────────────────────────────────────────
            elif s in ('---', '***', '___'):
                pass  # skip

            # ── Regular paragraph ─────────────────────────────────────────────
            elif s:
                doc.add_paragraph(s)

            i += 1

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc.save(tmp.name)
            data = pathlib.Path(tmp.name).read_bytes()

        # Clean up the mermaid temp dir
        try:
            _fallback_tmpdir_obj.cleanup()
        except Exception:
            pass

        # Write to disk cache for future requests
        try:
            cache_path.write_bytes(data)
        except Exception as cache_err:
            log.warning("docx.cache_write_failed", extra={"error": str(cache_err)})

        return StreamingResponse(
            iter([data]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
        )
    except ImportError:
        raise HTTPException(500, "pandoc not found and python-docx not installed.")
    except Exception as exc:
        log.exception("docx.conversion_failed", extra={"file": filename})
        raise HTTPException(500, f"Conversion failed: {exc}")


# ── Phase status (polling endpoint for UI) ─────────────────────────────────────

@app.get("/api/v1/projects/{project_id}/status", tags=["pipeline"])
async def get_project_status(project_id: int):
    """Lightweight status poll — returns phase_statuses without full conversation.

    A2.1 — also returns:
      - `requirements_hash`: current frozen SHA256 of the P1 lock (None if
        requirements have never been confirmed).
      - `requirements_frozen_at`: ISO-8601 timestamp.
      - `stale_phase_ids`: downstream phases whose last completion used an
        OLDER requirements hash than the current one — used by the UI to
        render "⚠ outdated" badges and enable the "Re-run stale" button.
    """
    svc = _project_svc()
    proj = svc.get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    stale = svc.get_stale_phase_ids(project_id)

    # D2.1 — Judge Mode verification summary. Pulled from the persisted
    # audit report / requirements lock when available so the frontend
    # overlay has a single endpoint to hit.
    audit_summary: dict[str, object] = {}
    cascade_summary: dict[str, object] = {}
    resolved_citation_count: int | None = None
    part_check_count: int | None = None
    try:
        import json as _json
        from pathlib import Path as _Path
        out_dir = _Path("output") / f"project_{project_id}"
        audit_json_path = out_dir / "audit_report.json"
        if audit_json_path.exists():
            rep = _json.loads(audit_json_path.read_text(encoding="utf-8"))
            issues = rep.get("issues", [])
            audit_summary = {
                "overall_pass": bool(rep.get("overall_pass")),
                "cascade_errors": sum(
                    1 for i in issues if i.get("category") == "cascade_vs_claims"
                ),
                "unresolved_citations": sum(
                    1 for i in issues if i.get("category") == "citations"
                ),
            }
            cascade_summary = {
                "computed_nf_db": rep.get("computed_nf_db"),
                "claimed_nf_db": rep.get("claimed_nf_db"),
                "computed_gain_db": rep.get("computed_gain_db"),
                "claimed_gain_db": rep.get("claimed_gain_db"),
            }
            resolved_citation_count = rep.get("resolved_citation_count")
            part_check_count = rep.get("part_check_count")
    except Exception as _e:  # noqa: BLE001 — surface of a *read-only* summary
        log.debug("Judge Mode summary load failed for project %s: %s",
                  project_id, _e)

    # Phase-level applicability flags — the frontend uses these to disable
    # the Run button and show the NOT APPLICABLE label, so the source of
    # truth is backend-derived rather than browser localStorage.
    from services.phase_scopes import PHASE_APPLICABLE_SCOPES
    design_scope = (proj.get("design_scope") or "full").lower()
    applicable_phases = sorted(
        pid for pid, scopes in PHASE_APPLICABLE_SCOPES.items()
        if design_scope in scopes
    )

    return {
        "project_id": project_id,
        "current_phase": proj.get("current_phase"),
        "design_scope": design_scope,
        "applicable_phase_ids": applicable_phases,
        "phase_statuses": proj.get("phase_statuses", {}),
        "requirements_hash": proj.get("requirements_hash"),
        "requirements_frozen_at": proj.get("requirements_frozen_at"),
        "stale_phase_ids": stale,
        "audit_summary": audit_summary,
        "cascade_summary": cascade_summary,
        "resolved_citation_count": resolved_citation_count,
        "part_check_count": part_check_count,
    }


# A2.2 — re-run all stale downstream phases (one-click recovery after the
# user edits requirements in P1).
@app.post("/api/v1/projects/{project_id}/pipeline/rerun-stale", tags=["pipeline"])
async def rerun_stale_phases(project_id: int, background_tasks: BackgroundTasks):
    """
    Reset every phase whose last completion used an older requirements_hash
    than the currently-locked one, and kick off the pipeline to re-run them.

    Returns a 200 with `reset_phases: []` when there is nothing stale. Never
    regenerates phases that were already pending/running.
    """
    svc = _project_svc()
    proj = svc.get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    stale = svc.get_stale_phase_ids(project_id)
    if not stale:
        return {
            "status": "no_stale_phases",
            "project_id": project_id,
            "reset_phases": [],
        }

    for phase_id in stale:
        svc.set_phase_status(project_id, phase_id, "pending")

    log.info(
        "api.rerun_stale",
        extra={"project_id": project_id, "stale_phases": stale},
    )

    background_tasks.add_task(_pipeline_svc().run_pipeline, project_id)
    return {
        "status": "pipeline_started",
        "project_id": project_id,
        "reset_phases": stale,
    }


# E2 — dry-run preview of the rerun plan the "Re-run stale" button will execute.
# Uses `services.stale_phases.rerun_plan` which adds manual-phase warnings and a
# human summary — safe to call any number of times, does not mutate DB state.
@app.get("/api/v1/projects/{project_id}/pipeline/rerun-plan", tags=["pipeline"])
async def get_rerun_plan(project_id: int):
    """
    Return the advisory plan for re-running stale phases.

    Payload shape:
        {
          "project_id": 1,
          "current_hash": "abc123…" | None,
          "stale": ["P2", "P4"],
          "order": ["P2", "P3", "P4", ...],
          "blocked_by_manual": ["P4"],     # manual phases (PCB/FPGA) that
                                            #  would also need rework
          "status_summary": {"P1": "fresh", "P2": "stale", ...},
          "summary": "2 stale phases — re-run in order P2 → P4"
        }

    The frontend RerunPlanDrawer uses this to render "will re-run" badges
    before the user confirms by hitting `/pipeline/rerun-stale`.
    """
    svc = _project_svc()
    proj = svc.get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")

    from services.stale_phases import rerun_plan, phase_status_summary

    plan = rerun_plan(proj)
    # `rerun_plan` treats the input as opaque row; it will read
    # requirements_hash + phase_statuses fine from the dict.
    summary = phase_status_summary(proj)
    plan["project_id"] = project_id
    plan["status_summary"] = summary
    return plan


# E1 — Judge-mode wipe-state: one-shot reset of mutable project fields so the
# panel can reboot to a clean demo state in under a second. Identity fields
# (name/description/design_type/output_dir) are preserved so the judges see
# the same project tile in the left panel.
@app.post("/api/v1/projects/{project_id}/reset-state", tags=["projects"])
async def reset_project_state(project_id: int):
    """
    Wipe a project's phase statuses, conversation history, design parameters
    and requirements lock columns. Does NOT delete the project row.

    Used by JudgeMode's "Clear state" button — the overlay then re-polls
    `/status` and everything displays the pre-demo baseline.
    """
    svc = _project_svc()
    proj = svc.get(project_id)
    if not proj:
        raise HTTPException(404, f"Project {project_id} not found")
    result = svc.reset_state(project_id)
    return {"status": "reset", **result}


# ── Test UI (standalone HTML workflow tester) ──────────────────────────────────

@app.get("/testui", response_class=HTMLResponse, tags=["ops"])
async def test_ui():
    """Serve standalone HTML workflow test page (no Streamlit needed)."""
    import pathlib
    p = pathlib.Path(__file__).parent / "test_ui.html"
    if p.exists():
        return HTMLResponse(content=p.read_text(), status_code=200)
    return HTMLResponse(content="<h1>test_ui.html not found</h1>", status_code=404)


@app.get("/app", response_class=HTMLResponse, tags=["ops"])
async def serve_frontend():
    """Serve the React v5 frontend bundle at http://localhost:8000/app"""
    import pathlib
    p = pathlib.Path(__file__).parent / "frontend" / "bundle.html"
    if p.exists():
        return HTMLResponse(
            content=p.read_text(encoding="utf-8", errors="replace"),
            status_code=200,
            headers={
                "Cache-Control": "no-store, no-cache, must-revalidate",
                "Pragma": "no-cache",
            },
        )
    return HTMLResponse(content="<h1>Frontend not built yet. Run the React build.</h1>", status_code=404)


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host=settings.fastapi_host,
        port=settings.fastapi_port,
        reload=settings.debug,
        log_level=settings.log_level.lower(),
    )